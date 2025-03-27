import openai
import os
import config
import pandas as pd
from docx import Document

# Set up OpenAI API
openai.api_key = config.CHATGPT_API_KEY

def evaluate_answers(excel_path):
    df = pd.read_excel(excel_path)

    feedback_list = []
    marks_list = []

    for index, row in df.iterrows():
        question = row["Question"]
        answer = row["Answer"]

        # Send request to ChatGPT for evaluation
        prompt = f"Evaluate the following answer:\n\nQuestion: {question}\nAnswer: {answer}\n\nProvide a score out of 10 and feedback."
        
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )

        result = response["choices"][0]["message"]["content"]
        feedback, marks = result.split("\n")

        feedback_list.append(feedback.strip())
        marks_list.append(marks.strip())

    df["Feedback"] = feedback_list
    df["Marks"] = marks_list

    # Save as DOCX
    doc_path = os.path.join(config.OUTPUT_FOLDER, "evaluation.docx")
    doc = Document()
    doc.add_heading("Evaluation Report", level=1)

    for i, row in df.iterrows():
        doc.add_paragraph(f"Question {row['Question No.']}: {row['Question']}")
        doc.add_paragraph(f"Answer: {row['Answer']}")
        doc.add_paragraph(f"Feedback: {row['Feedback']}")
        doc.add_paragraph(f"Marks: {row['Marks']}")
        doc.add_paragraph("\n---\n")

    doc.save(doc_path)

    return doc_path
